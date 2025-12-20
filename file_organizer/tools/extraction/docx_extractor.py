# from .base_extractor import BaseExtractor

# class DocxExtractor(BaseExtractor):
#     def extract(self, file_path: str) -> str:
#         """Extract text from Word documents."""
#         try:
#             from docx import Document
#             doc = Document(file_path)
#             return "\n".join(para.text for para in doc.paragraphs if para.text)
#         except ImportError:
#             raise ImportError("Install python-docx: pip install python-docx")


from .base_extractor import BaseExtractor
import logging

logger = logging.getLogger(__name__)

class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> dict:
        """
        Extract text from Word documents.
        
        Returns:
            dict: Dictionary containing 'content', 'success', and 'error' keys
        """
        try:
            from docx import Document
            doc = Document(file_path)
            content = "\n".join(para.text for para in doc.paragraphs if para.text)
            return {
                'content': content,
                'success': True,
                'file_path': file_path
            }
        except ImportError as e:
            error_msg = "Install python-docx: pip install python-docx"
            logger.error(error_msg)
            return {
                'content': '',
                'success': False,
                'error': error_msg,
                'file_path': file_path
            }
        except Exception as e:
            error_msg = f"Error extracting text from {file_path}: {str(e)}"
            logger.error(error_msg)
            return {
                'content': '',
                'success': False,
                'error': error_msg,
                'file_path': file_path
            }